#!/usr/bin/env python3
"""
УЛУЧШЕННЫЙ ГЕНЕРАТОР ДОКУМЕНТОВ v2.0
- Красивое форматирование таблиц с границами
- Поддержка транзитных адресов
- Точное соответствие оригинальному формату
- Выпадающие списки для выбора
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook
import pandas as pd
from datetime import datetime


class EnhancedDocumentGenerator:
    """Улучшенный генератор с красивым форматированием"""
    
    def __init__(self, input_file):
        self.input_file = input_file
        self.wb = load_workbook(input_file)
        self.tasks_data = []
        self.prices = {}
        self.load_data()
    
    def load_data(self):
        """Загрузка данных с поддержкой транзитных адресов"""
        # Загрузка прайс-листа
        prices_sheet = self.wb['расценки']
        for row in prices_sheet.iter_rows(min_row=2, values_only=True):
            if row[0] and row[1] and row[2]:
                self.prices[int(row[0])] = {
                    'description': row[1],
                    'price': float(row[2])
                }
        
        # Загрузка заданий с транзитными адресами
        df = pd.read_excel(self.input_file, sheet_name='Эксель')
        for idx, row in df.iterrows():
            if pd.notna(row.iloc[0]) and row.iloc[0] != 1:
                address = row.iloc[-1] if pd.notna(row.iloc[-1]) else ''
                
                # Разбираем транзитные адреса
                transit_addresses = []
                if 'транзит' in address.lower():
                    parts = address.split('транзитные адреса')
                    main_address = parts[0].strip()
                    if len(parts) > 1:
                        transit_text = parts[1].strip()
                        # Разделяем транзитные адреса по запятым
                        transit_addresses = [addr.strip() for addr in transit_text.split(',') if addr.strip()]
                else:
                    main_address = address
                
                task = {
                    'number': row.iloc[0],
                    'month': row.iloc[2] if pd.notna(row.iloc[2]) else '',
                    'district': row.iloc[3] if pd.notna(row.iloc[3]) else '',
                    'main_address': main_address,
                    'transit_addresses': transit_addresses,
                    'full_address': address,
                    'notes': row.iloc[4] if pd.notna(row.iloc[4]) else '',
                }
                self.tasks_data.append(task)
    
    def set_cell_border(self, cell, **kwargs):
        """
        Установка границ ячейки таблицы
        Параметры: top, bottom, left, right, insideH, insideV
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Создаем элемент границ
        tcBorders = OxmlElement('w:tcBorders')
        
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            if edge in kwargs:
                edge_data = kwargs.get(edge)
                edge_el = OxmlElement(f'w:{edge}')
                edge_el.set(qn('w:val'), edge_data.get('val', 'single'))
                edge_el.set(qn('w:sz'), str(edge_data.get('sz', 4)))
                edge_el.set(qn('w:space'), str(edge_data.get('space', 0)))
                edge_el.set(qn('w:color'), edge_data.get('color', '000000'))
                tcBorders.append(edge_el)
        
        tcPr.append(tcBorders)
    
    def set_cell_background(self, cell, color):
        """Установка цвета фона ячейки"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def generate_beautiful_zadanie(self, task_number, output_folder='output'):
        """Генерация КРАСИВО оформленного задания"""
        import os
        os.makedirs(output_folder, exist_ok=True)
        
        doc = Document()
        
        # === ШАПКА ДОКУМЕНТА ===
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run('Приложение № 1 к Договору на оказание услуг\n')
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        
        run = header.add_run('№ 01-4/25 от «10» апреля 2025 г.')
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.font.bold = True
        
        doc.add_paragraph()
        
        # Email
        email = doc.add_paragraph()
        email.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = email.add_run('На электронную почту: 9624294@gmail.com')
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        
        doc.add_paragraph()
        
        # === ЗАГОЛОВОК ===
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('ЗАДАНИЕ №\nна оказание услуг\nоб оказанных услугах')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = 'Arial'
        
        # Номер задания
        title2 = doc.add_paragraph()
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title2.add_run(f'\n{task_number}')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = 'Arial'
        
        # Город и дата
        city = doc.add_paragraph()
        city.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = city.add_run(f'г. Санкт-Петербург')
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        
        city.add_run(' ' * 50)  # Пробелы
        
        run = city.add_run(datetime.now().strftime("%d.%m.%Y"))
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # ФИО
        fio = doc.add_paragraph()
        run = fio.add_run('Ф.И.О., должность лица, направившего задание – Багров М.С.')
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        
        # Примечание
        note = doc.add_paragraph()
        run = note.add_run('Данные по заданию (в случае отправки Исполнителю по электронной почте).')
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.italic = True
        
        doc.add_paragraph()
        
        # === ТАБЛИЦА С КРАСИВЫМ ОФОРМЛЕНИЕМ ===
        
        # Определяем количество строк
        tasks_for_number = [t for t in self.tasks_data if str(t['number']) == str(task_number)]
        num_rows = sum(1 + len(t.get('transit_addresses', [])) for t in tasks_for_number)
        
        table = doc.add_table(rows=num_rows + 1, cols=4)
        table.style = 'Table Grid'
        
        # Настройка ширины столбцов
        table.autofit = False
        table.allow_autofit = False
        widths = [Inches(3), Inches(1.5), Inches(1.5), Inches(3)]
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = width
        
        # === ЗАГОЛОВКИ ТАБЛИЦЫ ===
        headers = ['Адрес\nпредоставления услуги', 'Дата передачи\nзадания', 'Дата\nвыполнения задания', 'Вид оказанной услуги']
        
        for i, header_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header_text
            
            # Форматирование заголовка
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
            
            # Серый фон для заголовка
            self.set_cell_background(cell, 'D9D9D9')
            
            # Границы
            border_settings = {
                'top': {'val': 'single', 'sz': 6, 'color': '000000'},
                'bottom': {'val': 'single', 'sz': 6, 'color': '000000'},
                'left': {'val': 'single', 'sz': 6, 'color': '000000'},
                'right': {'val': 'single', 'sz': 6, 'color': '000000'}
            }
            self.set_cell_border(cell, **border_settings)
        
        # === ДАННЫЕ ТАБЛИЦЫ ===
        row_idx = 1
        for task in tasks_for_number:
            # Основной адрес
            cells = table.rows[row_idx].cells
            
            # Адрес
            cells[0].text = task['main_address']
            
            # Дата передачи
            cells[1].text = datetime.now().strftime('%d.%m.%Y')
            
            # Дата выполнения (пусто для заполнения)
            cells[2].text = ''
            
            # Вид услуги
            service_text = task['notes'][:200] if task['notes'] else ''
            cells[3].text = service_text
            
            # Форматирование строки данных
            for cell in cells:
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.TOP
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                
                # Границы
                border_settings = {
                    'top': {'val': 'single', 'sz': 4, 'color': '000000'},
                    'bottom': {'val': 'single', 'sz': 4, 'color': '000000'},
                    'left': {'val': 'single', 'sz': 4, 'color': '000000'},
                    'right': {'val': 'single', 'sz': 4, 'color': '000000'}
                }
                self.set_cell_border(cell, **border_settings)
            
            row_idx += 1
            
            # Транзитные адреса (если есть)
            if task.get('transit_addresses'):
                for transit_addr in task['transit_addresses']:
                    cells = table.rows[row_idx].cells
                    
                    cells[0].text = f"   → {transit_addr}"  # С отступом
                    cells[1].text = datetime.now().strftime('%d.%m.%Y')
                    cells[2].text = ''
                    cells[3].text = '(транзитный адрес)'
                    
                    # Форматирование транзитных адресов
                    for cell in cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                                run.font.name = 'Arial'
                                run.font.italic = True
                                run.font.color.rgb = RGBColor(128, 128, 128)
                        
                        # Светло-серый фон для транзитных
                        self.set_cell_background(cell, 'F2F2F2')
                        
                        border_settings = {
                            'top': {'val': 'single', 'sz': 4, 'color': 'CCCCCC'},
                            'bottom': {'val': 'single', 'sz': 4, 'color': 'CCCCCC'},
                            'left': {'val': 'single', 'sz': 4, 'color': 'CCCCCC'},
                            'right': {'val': 'single', 'sz': 4, 'color': 'CCCCCC'}
                        }
                        self.set_cell_border(cell, **border_settings)
                    
                    row_idx += 1
        
        # Подписи
        doc.add_paragraph()
        doc.add_paragraph()
        
        signature_table = doc.add_table(rows=2, cols=2)
        signature_table.style = 'Table Grid'
        
        # Заказчик
        signature_table.rows[0].cells[0].text = 'Заказчик:'
        signature_table.rows[1].cells[0].text = '_______________ (подпись)'
        
        # Исполнитель
        signature_table.rows[0].cells[1].text = 'Исполнитель:'
        signature_table.rows[1].cells[1].text = '_______________ (подпись)'
        
        # Сохранение
        filename = f'{output_folder}/Задание_{task_number}_КРАСИВОЕ.docx'
        doc.save(filename)
        return filename


# Пример использования
if __name__ == '__main__':
    generator = EnhancedDocumentGenerator('/mnt/user-data/uploads/ПРИЛОЖЕНИЕ_1_3_4_ALL_IN_ONE_Задание_Отчет_АКТ_месяц_ноябрь_2025.xlsx')
    
    # Генерация с красивым форматированием
    filename = generator.generate_beautiful_zadanie('11-1')
    print(f"Создан документ: {filename}")
    
    # Показываем данные о транзитных адресах
    print("\n=== НАЙДЕННЫЕ ТРАНЗИТНЫЕ АДРЕСА ===")
    for task in generator.tasks_data:
        if task.get('transit_addresses'):
            print(f"\nЗадание {task['number']}:")
            print(f"  Основной: {task['main_address']}")
            print(f"  Транзитных: {len(task['transit_addresses'])}")
            for addr in task['transit_addresses']:
                print(f"    → {addr}")
