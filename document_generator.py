#!/usr/bin/env python3
"""
Система автоматической генерации документов для оказания услуг
Автоматически создает: Задания, Отчеты и Акты на основе данных из Excel
"""

import pandas as pd
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json
import os


class DocumentGenerator:
    """Генератор документов для договоров об оказании услуг"""
    
    def __init__(self, input_file):
        self.input_file = input_file
        self.wb = load_workbook(input_file)
        self.tasks_data = []
        self.prices = {}
        self.load_data()
    
    def load_data(self):
        """Загрузка данных из Excel файла"""
        # Загрузка прайс-листа
        prices_sheet = self.wb['расценки']
        for row in prices_sheet.iter_rows(min_row=2, values_only=True):
            if row[0] and row[1] and row[2]:
                self.prices[int(row[0])] = {
                    'description': row[1],
                    'price': float(row[2])
                }
        
        # Загрузка заданий
        df = pd.read_excel(self.input_file, sheet_name='Эксель')
        for idx, row in df.iterrows():
            if pd.notna(row.iloc[0]) and row.iloc[0] != 1:
                task = {
                    'number': row.iloc[0],
                    'month': row.iloc[2] if pd.notna(row.iloc[2]) else '',
                    'district': row.iloc[3] if pd.notna(row.iloc[3]) else '',
                    'address': row.iloc[-1] if pd.notna(row.iloc[-1]) else '',
                    'notes': row.iloc[4] if pd.notna(row.iloc[4]) else '',
                }
                self.tasks_data.append(task)
    
    def generate_zadanie(self, task_number, output_folder='output'):
        """Генерация документа Задание"""
        os.makedirs(output_folder, exist_ok=True)
        
        doc = Document()
        
        # Настройка документа
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        
        # Заголовок
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run('Приложение № 1 к Договору на оказание услуг\n')
        run = header.add_run('№ 01-4/25 от «10» апреля 2025 г.')
        run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Email
        email = doc.add_paragraph()
        email.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        email.add_run('На электронную почту: 9624294@gmail.com')
        
        doc.add_paragraph()
        
        # Название документа
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(f'ЗАДАНИЕ №\nна оказание услуг\nоб оказанных услугах\n\n{task_number}')
        title_run.font.size = Pt(14)
        title_run.bold = True
        
        # Город и дата
        city = doc.add_paragraph()
        city.alignment = WD_ALIGN_PARAGRAPH.CENTER
        city.add_run(f'г. Санкт-Петербург                                    {datetime.now().strftime("%d.%m.%Y")}')
        
        doc.add_paragraph()
        
        # ФИО
        doc.add_paragraph('Ф.И.О., должность лица, направившего задание – Багров М.С.')
        
        # Данные по заданию
        doc.add_paragraph('Данные по заданию (в случае отправки Исполнителю по электронной почте).')
        
        doc.add_paragraph()
        
        # Таблица заданий
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Заголовки таблицы
        headers = ['Адрес\nпредоставления услуги', 'Дата передачи задания', 'Дата\nвыполнения задания', 'Вид оказанной услуги']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Добавляем строки с данными
        for task in self.tasks_data:
            if str(task['number']) == str(task_number):
                row = table.add_row()
                row.cells[0].text = task['address']
                row.cells[1].text = datetime.now().strftime('%d.%m.%Y')
                row.cells[2].text = ''
                row.cells[3].text = task['notes'][:100] if task['notes'] else ''
        
        # Сохранение
        filename = f'{output_folder}/Задание_{task_number}.docx'
        doc.save(filename)
        return filename
    
    def generate_otchet(self, task_number, services, output_folder='output'):
        """Генерация документа Отчет"""
        os.makedirs(output_folder, exist_ok=True)
        
        doc = Document()
        
        # Настройка документа
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        
        # Заголовок
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run('Приложение № 3 к Договору на оказание услуг\n')
        run = header.add_run('№ 01-4/25 от «10» апреля 2025 г.')
        run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Адресат
        addressee = doc.add_paragraph()
        addressee.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        addressee.add_run('Генеральному директору АО «Северен-Телеком»\n')
        addressee.add_run('Ковтонюку А.В.')
        
        doc.add_paragraph()
        
        # Название документа
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(f'ОТЧЕТ №\nоб оказанных услугах\n\n{task_number}')
        title_run.font.size = Pt(14)
        title_run.bold = True
        
        # Город и дата
        city = doc.add_paragraph()
        city.alignment = WD_ALIGN_PARAGRAPH.CENTER
        today = datetime.now().strftime('%d.%m.%Y')
        city.add_run(f'г. Санкт-Петербург                                    {today}')
        
        doc.add_paragraph()
        
        # Текст отчета
        doc.add_paragraph(f'По состоянию на {today} г. в целях оказания услуг и полученных заданий к Договору об оказании услуг № 01-4/25 от «10» апреля 2025 г. Исполнитель оказал следующие услуги:')
        
        doc.add_paragraph()
        
        # Таблица услуг
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Заголовки
        headers = ['Адрес\nпредоставления услуги', 'Дата передачи задания', 'Дата\nвыполнения задания', 'Вид оказанной услуги', 'Стоимость оказанных услуг, в руб.']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        total = 0
        for service in services:
            row = table.add_row()
            task = next((t for t in self.tasks_data if str(t['number']) == str(task_number)), None)
            
            if task:
                row.cells[0].text = task['address']
            row.cells[1].text = service.get('start_date', datetime.now().strftime('%d.%m.%Y'))
            row.cells[2].text = service.get('end_date', datetime.now().strftime('%d.%m.%Y'))
            
            service_type = service.get('type', 1)
            if service_type in self.prices:
                desc = self.prices[service_type]['description']
                price = self.prices[service_type]['price']
                row.cells[3].text = desc
                row.cells[4].text = f"{price:,.2f}"
                total += price
        
        # Итого
        doc.add_paragraph()
        total_p = doc.add_paragraph()
        total_p.add_run(f'Итого: {total:,.2f} руб.').bold = True
        
        # Сохранение
        filename = f'{output_folder}/Отчет_{task_number}.docx'
        doc.save(filename)
        return filename
    
    def generate_akt(self, task_number, services, output_folder='output'):
        """Генерация документа Акт выполненных работ"""
        os.makedirs(output_folder, exist_ok=True)
        
        doc = Document()
        
        # Настройка документа
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        
        # Заголовок
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_run('Приложение № 4\n')
        header.add_run('к Договору об оказании услуг\n')
        header.add_run('№ 01-4/25 от «10» апреля 2025 г.')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Название документа
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(f'АКТ №\n{task_number}')
        title_run.font.size = Pt(14)
        title_run.bold = True
        
        doc.add_paragraph()
        
        # Город и дата
        city = doc.add_paragraph()
        city.alignment = WD_ALIGN_PARAGRAPH.LEFT
        today = datetime.now().strftime('%d.%m.%Y')
        city.add_run(f'г. Санкт-Петербург                                    {today} г.')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Стороны договора
        parties = doc.add_paragraph()
        parties.add_run(
            'Акционерное общество «Северен-Телеком», далее именуемое «Заказчик», в лице генерального директора '
            'Ковтонюка Андрея Владимировича, действующего на основании устава, с одной стороны и '
            'Индивидуальный предприниматель Сагуров Андрей Петрович далее именуемое «Исполнитель», '
            'в лице Сагурова Андрея Петровича, действующего на основании паспорта 4123 №406829, '
            'выданного 17.10.2023 г. ГУ МВД России по г. Санкт-Петербургу и Ленинградской области, '
            'зарегистрированного по адресу: Ленинградская область, Всеволожский р-н, г. Всеволожск, '
            'ул. Центральная д.7 кв.32 в качестве индивидуального предпринимателя в Межрайонной инспекции '
            'Федеральной налоговой службе №2 по Ленинградской области (ОГРНИП 325470400039051) с другой стороны, '
            'далее совместно именуемые «Стороны», составили настоящий акт (далее - Акт) о нижеследующем:'
        )
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Расчет суммы
        total = sum(self.prices[s['type']]['price'] for s in services if s['type'] in self.prices)
        
        # Пункт 1
        p1 = doc.add_paragraph()
        p1.add_run(f'1. Исполнителем за период с {services[0].get("start_date", "01.11.2025")} г. по {services[-1].get("end_date", today)} г. ')
        p1.add_run('в рамках Договора об оказании услуг № 01-4/25 от «10» апреля 2025 г (далее - Договор) ')
        p1.add_run(f'были оказаны следующие услуги: в соответствии с отчетом № {task_number}.')
        
        # Пункт 2
        p2 = doc.add_paragraph()
        p2.add_run(f'2. Итого услуг по отчету оказано на сумму:')
        
        # Сумма прописью
        total_text = self.number_to_words(total)
        amount = doc.add_paragraph()
        amount_run = amount.add_run(f'{total:,.2f} ({total_text}) рублей, 00 копеек,')
        amount_run.bold = True
        
        # НДС
        doc.add_paragraph('НДС не облагается в связи с применением Исполнителем упрощенной системы налогообложения.')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Подписи
        signatures = doc.add_table(rows=3, cols=2)
        signatures.cell(0, 0).text = 'Заказчик:'
        signatures.cell(0, 1).text = 'Исполнитель:'
        signatures.cell(1, 0).text = 'АО «Северен-Телеком»'
        signatures.cell(1, 1).text = 'ИП Сагуров А.П.'
        signatures.cell(2, 0).text = '_______________ Ковтонюк А.В.'
        signatures.cell(2, 1).text = '_______________ Сагуров А.П.'
        
        # Сохранение
        filename = f'{output_folder}/Акт_{task_number}.docx'
        doc.save(filename)
        return filename
    
    @staticmethod
    def number_to_words(number):
        """Преобразование числа в текст (упрощенная версия)"""
        thousands = int(number // 1000)
        hundreds = int(number % 1000)
        
        result = []
        
        if thousands > 0:
            ones = ['', 'одна', 'две', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
            tens = ['', 'десять', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
            
            if thousands >= 10:
                result.append(tens[thousands // 10])
            if thousands % 10 > 0:
                result.append(ones[thousands % 10])
            
            if thousands % 10 == 1 and thousands % 100 != 11:
                result.append('тысяча')
            elif thousands % 10 in [2, 3, 4] and thousands % 100 not in [12, 13, 14]:
                result.append('тысячи')
            else:
                result.append('тысяч')
        
        if hundreds > 0:
            hundreds_words = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот', 'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']
            tens_words = ['', 'десять', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
            ones_words = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
            
            result.append(hundreds_words[hundreds // 100])
            result.append(tens_words[(hundreds % 100) // 10])
            result.append(ones_words[hundreds % 10])
        
        return ' '.join(filter(None, result))
    
    def generate_all_documents(self, task_number, services, output_folder='output'):
        """Генерация всех документов для задания"""
        files = {
            'zadanie': self.generate_zadanie(task_number, output_folder),
            'otchet': self.generate_otchet(task_number, services, output_folder),
            'akt': self.generate_akt(task_number, services, output_folder)
        }
        return files


if __name__ == '__main__':
    # Пример использования
    generator = DocumentGenerator('/mnt/user-data/uploads/ПРИЛОЖЕНИЕ_1_3_4_ALL_IN_ONE_Задание_Отчет_АКТ_месяц_ноябрь_2025.xlsx')
    
    # Пример генерации документов для задания 11-1
    task_number = '11-1'
    services = [
        {'type': 1, 'start_date': '01.11.2025', 'end_date': '03.12.2025'},
        {'type': 1, 'start_date': '01.11.2025', 'end_date': '03.12.2025'},
        {'type': 4, 'start_date': '01.11.2025', 'end_date': '03.12.2025'},
    ]
    
    files = generator.generate_all_documents(task_number, services)
    
    print("Сгенерированные файлы:")
    for doc_type, filepath in files.items():
        print(f"  {doc_type}: {filepath}")
